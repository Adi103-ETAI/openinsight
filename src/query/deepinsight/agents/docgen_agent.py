# BUILT: docgen_agent
from __future__ import annotations

import asyncio
import json
import os
from typing import Any, Dict, Optional
from datetime import datetime
from pathlib import Path

from src.config.settings import get_settings
from src.services.llm.router import LLMRouter
from src.query.deepinsight.agents.skills import get_system_prompt


class DocGenResult:
    """Result of document generation processing."""
    
    def __init__(
        self, 
        file_path: str, 
        format: str, 
        page_count: int, 
        size_bytes: int, 
        title: str, 
        generated_at: str
    ):
        self.file_path = file_path
        self.format = format
        self.page_count = page_count
        self.size_bytes = size_bytes
        self.title = title
        self.generated_at = generated_at


class DocGenAgent:
    """
    Document generation agent for DeepInsight.
    
    Converts the final synthesized clinical answer into formatted DOCX or PDF file.
    Does NOT regenerate content — formats existing synthesis output.
    Triggered when doc_format is set in orchestrator routing decision.
    """
    
    def __init__(self, settings: Any, llm_router: LLMRouter):
        self.settings = settings
        self.llm_router = llm_router
        self.output_dir = Path("/tmp") / "openinsight_reports"
        self.output_dir.mkdir(exist_ok=True)
        
    async def run(
        self,
        doc_request: Dict[str, Any]
    ) -> DocGenResult:
        """
        Generate a document from synthesis content.
        
        Args:
            doc_request: {
                format: "pdf" | "docx",
                title: str,
                content: str,
                citations: List[Dict],
                patient_context: Optional[str],
                generated_at: str
            }
            
        Returns:
            DocGenResult with file path and metadata
        """
        doc_format = doc_request.get("format", "pdf")
        title = doc_request.get("title", "Clinical Summary")
        content = doc_request.get("content", "")
        citations = doc_request.get("citations", [])
        patient_context = doc_request.get("patient_context")
        generated_at = doc_request.get("generated_at", datetime.now().isoformat())
        
        if not content.strip():
            raise ValueError("Cannot generate document from empty content")
            
        # Get LLM client for document generation
        client = self.llm_router.get_client_for_agent("docgen")
        
        # Build document generation prompt from skill
        system_prompt = get_system_prompt("docgen_agent")
        
        context = f"""
DOC_REQUEST:
{json.dumps(doc_request, indent=2)}
"""
        
        # Call LLM for document structure assembly
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": context}
        ]
        
        response = await client.chat_completions(messages=messages)
        assembly_result = response.choices[0].message.content
        
        # Parse structured response (expected JSON format)
        try:
            doc_data = json.loads(assembly_result)
            sections = doc_data.get("sections", {})
            
        except (json.JSONDecodeError, KeyError):
            # Fallback: build sections manually
            sections = self._build_sections_fallback(
                content, citations, title, patient_context, generated_at
            )
        
        # Generate actual document file
        if doc_format == "docx":
            file_path = self._generate_docx(sections, title)
        elif doc_format == "pdf":
            file_path = self._generate_pdf(sections, title)
        else:
            raise ValueError(f"Unsupported document format: {doc_format}")
        
        # Get file metadata
        file_size = os.path.getsize(file_path)
        page_count = self._count_pages(file_path)
        
        return DocGenResult(
            file_path=str(file_path),
            format=doc_format,
            page_count=page_count,
            size_bytes=file_size,
            title=title,
            generated_at=generated_at
        )
    
    def _build_sections_fallback(
        self,
        content: str,
        citations: List[Dict],
        title: str,
        patient_context: Optional[str],
        generated_at: str
    ) -> Dict[str, str]:
        """Build document sections manually when LLM parsing fails."""
        
        # Build citations section
        citations_list = []
        for citation in citations:
            if citation.get("source_type") == "corpus":
                citations_list.append(
                    f"  {citation.get('claim_id', '')} — {citation.get('source_title', '')}. "
                    f"{citation.get('source_url', '')}"
                )
            else:  # web
                citations_list.append(
                    f"  {citation.get('claim_id', '')} — {citation.get('source_title', '')}. "
                    f"{citation.get('source_url', '')}"
                )
        
        sections = {
            "header": f"OpenInsight Clinical Summary\nGenerated: {generated_at}\nQuery: {title}",
            "clinical_summary": content,
            "sources_and_citations": "SOURCES AND CITATIONS:\n" + "\n".join(citations_list),
            "disclaimer": (
                "This report was generated by OpenInsight, an AI-powered clinical "
                "decision support system developed by SentArc Labs. It is intended "
                "to support, not replace, clinical judgment. Verify all drug dosages "
                "and guidelines against current institutional protocols before "
                f"prescribing. Generated: {generated_at}."
            )
        }
        
        if patient_context:
            sections["patient_context"] = f"Patient Context:\n{patient_context}"
            
        return sections
    
    def _generate_docx(self, sections: Dict[str, str], title: str) -> Path:
        """Generate DOCX document using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Header
            if "header" in sections:
                p = doc.add_paragraph(sections["header"])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("━" * 50)
                doc.add_paragraph()
            
            # Clinical Summary
            if "clinical_summary" in sections:
                heading = doc.add_heading("Clinical Summary", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph(sections["clinical_summary"])
            
            # Patient Context
            if "patient_context" in sections:
                heading = doc.add_heading("Patient Context", level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph(sections["patient_context"])
            
            # Sources and Citations
            if "sources_and_citations" in sections:
                heading = doc.add_heading("Sources and Citations", level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph(sections["sources_and_citations"])
            
            # Disclaimer
            if "disclaimer" in sections:
                p = doc.add_paragraph(sections["disclaimer"])
                p.style = 'Footer'
            
            # Save file
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"openinsight_report_{timestamp}.docx"
            file_path = self.output_dir / filename
            
            doc.save(file_path)
            return file_path
            
        except ImportError:
            raise ImportError("python-docx not installed for DOCX generation")
    
    def _generate_pdf(self, sections: Dict[str, str], title: str) -> Path:
        """Generate PDF document using reportlab."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Set up document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"openinsight_report_{timestamp}.pdf"
            file_path = self.output_dir / filename
            
            doc = SimpleDocTemplate(
                str(file_path),
                pagesize=A4,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch
            )
            
            # Custom styles
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                textColor=colors.darkblue,
                alignment=1  # centered
            ))
            
            styles.add(ParagraphStyle(
                name='CustomHeading2',
                parent=styles['Heading2'],
                fontSize=13,
                spaceAfter=12,
                textColor=colors.darkgray,
                spaceBefore=20
            ))
            
            styles.add(ParagraphStyle(
                name='CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                leading=14
            ))
            
            # Build content
            story = []
            
            # Header
            if "header" in sections:
                story.append(Paragraph(sections["header"], styles['CustomTitle']))
                story.append(Spacer(1, 12))
                story.append(Paragraph("━" * 50, styles['CustomBody']))
                story.append(Spacer(1, 20))
            
            # Clinical Summary
            if "clinical_summary" in sections:
                story.append(Paragraph("Clinical Summary", styles['CustomHeading2']))
                story.append(Paragraph(sections["clinical_summary"], styles['CustomBody']))
                story.append(Spacer(1, 12))
            
            # Patient Context
            if "patient_context" in sections:
                story.append(Paragraph("Patient Context", styles['CustomHeading2']))
                story.append(Paragraph(sections["patient_context"], styles['CustomBody']))
                story.append(Spacer(1, 12))
            
            # Sources and Citations
            if "sources_and_citations" in sections:
                story.append(Paragraph("Sources and Citations", styles['CustomHeading2']))
                story.append(Paragraph(sections["sources_and_citations"], styles['CustomBody']))
                story.append(Spacer(1, 12))
            
            # Disclaimer
            if "disclaimer" in sections:
                story.append(Paragraph(sections["disclaimer"], styles['CustomBody']))
            
            # Build PDF
            doc.build(story)
            return file_path
            
        except ImportError:
            raise ImportError("reportlab not installed for PDF generation")
    
    def _count_pages(self, file_path: Path) -> int:
        """Count pages in generated document (simplified)."""
        if file_path.suffix.lower() == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return len(reader.pages)
            except ImportError:
                return 1  # fallback
        else:
            return 1  # DOCX pages are complex to count without full library