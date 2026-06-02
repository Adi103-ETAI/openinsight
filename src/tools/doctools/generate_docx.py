"""generate_docx — render a structured dict of sections into a DOCX file."""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Dict

logger = logging.getLogger(__name__)


def generate_docx(sections: Dict[str, str], title: str, output_dir: Path = None) -> str:
    """
    Render sections into a DOCX using python-docx. Returns the file path.
    Returns "" if python-docx is not installed or an error occurs.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        from src.tools.doctools.generate_filename import generate_filename

        if output_dir is None:
            output_dir = Path("/tmp") / "openinsight_reports"
        output_dir.mkdir(parents=True, exist_ok=True)
        path = output_dir / generate_filename(title, "docx")

        doc = Document()
        for key, text in sections.items():
            if key == "header":
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif key == "disclaimer":
                doc.add_paragraph(text).style = "Footer"
            else:
                heading = key.replace("_", " ").title()
                doc.add_heading(heading, level=1)
                doc.add_paragraph(text)

        doc.save(str(path))
        return str(path)
    except ImportError:
        logger.warning("python-docx not installed; cannot generate DOCX")
        return ""
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        return ""
